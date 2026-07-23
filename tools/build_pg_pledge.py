from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "submission"
OUT_PATH = OUT_DIR / "벨로르_고가상품_판매_확약서.docx"

FONT = "Malgun Gothic"
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
NAVY = RGBColor(0x17, 0x2A, 0x46)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
MUTED = RGBColor(0x62, 0x69, 0x73)
LIGHT = "F2F4F7"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_font(run, size=11, bold=False, color=BLACK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_clause(doc, label, text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.1
    label_run = paragraph.add_run(f"{label}  ")
    set_font(label_run, bold=True, color=NAVY)
    text_run = paragraph.add_run(text)
    set_font(text_run)
    return paragraph


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, NAVY),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run("BELLORE  |  PG·카드사 심사 제출용")
    set_font(header_run, size=9, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer_run = footer.add_run("벨로르  |  bellore.co.kr  |  010-6293-6668")
    set_font(footer_run, size=8.5, color=MUTED)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(4)
    kicker_run = kicker.add_run("확약서")
    set_font(kicker_run, size=10, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run("500만원 이상 고가상품 판매 확약서")
    set_font(title_run, size=23, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.keep_with_next = True
    subtitle_run = subtitle.add_run("수신: 포트원 및 카드사·PG사 입점심사 담당자 귀중")
    set_font(subtitle_run, size=11, bold=True, color=MUTED)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = Pt(12)
    lead.paragraph_format.line_spacing = 1.1
    lead_run = lead.add_run(
        "당사는 500만원 이상 고가 명품시계를 판매하는 가맹점으로서 건전한 카드거래와 "
        "소비자 보호를 위하여 아래 사항을 확인하고 이를 성실히 준수할 것을 확약합니다."
    )
    set_font(lead_run, bold=True)

    metadata = doc.add_table(rows=5, cols=2)
    metadata.style = "Table Grid"
    metadata_rows = (
        ("상호", "벨로르"),
        ("대표자", "정성호"),
        ("사업자등록번호", "707-69-00729"),
        ("사업장 주소", "서울특별시 중구 다산로 258 벨로르"),
        ("쇼핑몰 / 연락처", "https://bellore.co.kr  /  010-6293-6668"),
    )
    for row, (label, value) in zip(metadata.rows, metadata_rows):
        left, right = row.cells
        shade_cell(left, LIGHT)
        left_p = left.paragraphs[0]
        left_p.paragraph_format.space_after = Pt(0)
        left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(left_p.add_run(label), size=10.5, bold=True, color=NAVY)
        right_p = right.paragraphs[0]
        right_p.paragraph_format.space_after = Pt(0)
        set_font(right_p.add_run(value), size=10.5)
    set_table_geometry(metadata, (2700, 6660))

    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("확약 내용")
    set_keep_with_next(heading)

    clauses = (
        (
            "제1조",
            "실물 명품시계의 정상적인 통신판매 거래에 대해서만 결제를 받으며, 현금융통·카드깡·"
            "본인 미사용 거래·상품권 및 가상자산 구입 등 변칙적인 자금융통 목적의 거래를 하지 않습니다.",
        ),
        (
            "제2조",
            "구매자 본인 명의 또는 적법한 사용 승낙을 받은 결제수단만 사용하도록 하며, 타인의 카드·"
            "계정·개인정보를 도용한 결제를 허용하지 않습니다.",
        ),
        (
            "제3조",
            "위조 상품, 허위 상품, 실제 판매가와 현저히 다른 가격의 상품 또는 배송 의사가 없는 상품을 "
            "판매하거나 결제하지 않습니다.",
        ),
        (
            "제4조",
            "고가 주문은 구매자 정보, 연락처, 거래 의사와 부정거래 징후를 확인하고, 의심 거래는 결제 확정과 "
            "출고를 보류하거나 취소합니다.",
        ),
        (
            "제5조",
            "현금융통, 명의도용, 부정거래 등 이상 거래가 확인되면 상품과 거래를 즉시 중단하고 결제 취소·환불을 "
            "처리하며, 카드사·PG사·포트원 및 관계기관의 확인 절차에 적극 협조합니다.",
        ),
        (
            "제6조",
            "bellore.co.kr에서 체결되는 모든 구매계약의 판매 당사자는 벨로르입니다. 벨로르는 상품정보 제공, "
            "주문확인, 결제, 검수, 배송, 청약철회, 교환·환불, 고객상담 및 거래분쟁 해결에 관한 책임을 부담합니다.",
        ),
        (
            "제7조",
            "상품 공급협력사는 구매자에 대한 판매자가 아니며, 공급 상품으로 인한 민원·손해·분쟁도 벨로르가 "
            "직접 접수하고 해결합니다.",
        ),
        (
            "제8조",
            "카드사·PG사·포트원의 운영정책과 관계 법령을 준수하고, 거래증빙·본인확인·배송증빙·정품 검수 기록을 "
            "보관하며 요청 시 제출합니다.",
        ),
        (
            "제9조",
            "위 확약을 위반하여 발생하는 결제 취소, 환불, 민원, 손해, 제재 및 계약 해지에 관한 책임을 부담하며, "
            "해당 이슈 발생 시 하위몰 또는 관련 서비스의 즉시 해지·중단 조치에 이의 없이 협조합니다.",
        ),
    )
    for label, text in clauses:
        add_clause(doc, label, text)

    acknowledgment = doc.add_paragraph()
    acknowledgment.paragraph_format.space_before = Pt(8)
    acknowledgment.paragraph_format.space_after = Pt(14)
    acknowledgment.paragraph_format.keep_with_next = True
    ack_run = acknowledgment.add_run(
        "본 확약서의 내용을 충분히 확인하였으며, 상기 사항을 성실히 준수할 것을 확약합니다."
    )
    set_font(ack_run, bold=True, color=NAVY)

    signature = doc.add_table(rows=4, cols=2)
    signature.style = "Table Grid"
    sig_rows = (
        ("작성일", "2026년 07월 23일"),
        ("상호", "벨로르"),
        ("대표자", "정성호  (서명 또는 인감)"),
        ("인감 날인란", "법인인감 또는 개인인감을 이 칸에 날인해 주세요."),
    )
    for row_index, (label, value) in enumerate(sig_rows):
        left, right = signature.rows[row_index].cells
        shade_cell(left, LIGHT)
        left_p = left.paragraphs[0]
        left_p.paragraph_format.space_after = Pt(0)
        left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(left_p.add_run(label), size=10.5, bold=True, color=NAVY)
        right_p = right.paragraphs[0]
        right_p.paragraph_format.space_after = Pt(0)
        if row_index == 3:
            right_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(right_p.add_run(value), size=9.5, color=MUTED)
            right.height = Inches(0.8)
        else:
            set_font(right_p.add_run(value), size=10.5)
    set_table_geometry(signature, (2700, 6660))

    doc.core_properties.title = "벨로르 500만원 이상 고가상품 판매 확약서"
    doc.core_properties.subject = "PG사 및 카드사 입점 심사 제출용"
    doc.core_properties.author = "벨로르"
    doc.core_properties.keywords = "벨로르, 고가상품, 판매 확약서, PG 심사"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
